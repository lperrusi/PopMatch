"""
Generates PopMatch Interview Preparation Guide as DOCX and PDF.
Run: python3 docs/generate_interview_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__))

# ─── Content Definition ──────────────────────────────────────────────────────

SECTIONS = [
    {
        "heading": "1. Project Overview — The 30-Second Pitch",
        "level": 1,
        "content": [
            ("p", "PopMatch is a cross-platform mobile app (Flutter/Dart) that uses a 7-service hybrid AI recommendation engine to surface personalized movies and TV shows through a Tinder-style swipe interface. Users swipe right (like), left (dislike), up (match + watchlist), or down (skip). Every interaction feeds seven concurrent ML systems that continuously refine future recommendations."),
            ("p", "Tech Stack:"),
            ("bullet", "Frontend: Flutter/Dart (iOS + Android from a single codebase)"),
            ("bullet", "Backend: Firebase Auth, Firestore, Cloud Functions (Node.js)"),
            ("bullet", "APIs: TMDB (movie catalog), OMDb (external ratings — IMDb, Rotten Tomatoes, Metacritic)"),
            ("bullet", "ML Persistence: SharedPreferences (local, cross-session ML state)"),
            ("bullet", "Architecture: Provider pattern (6 ChangeNotifiers), singleton services, 64-dimensional vector embeddings"),
            ("bullet", "AI/ML: Embeddings, Collaborative Filtering, Matrix Factorization, Online Learning, Adaptive Weighting, Behavioral Tracking, Contextual Bandits"),
        ],
    },
    {
        "heading": "2. Core AI Architecture — The Orchestrator Pattern",
        "level": 1,
        "content": [
            ("p", "The recommendation engine is a hybrid orchestrated system with 7 concurrent AI/ML services — architecturally identical to how modern AI agent systems work. The central method MovieProvider.loadPersonalizedRecommendations() acts as the ORCHESTRATOR (equivalent to an AI agent runner), coordinating specialized modules like tools."),
            ("h2", "Stage 1 — Retrieval (4 Parallel Discovery Strategies)"),
            ("bullet", "Trending movies from TMDB (real-time popularity signal)"),
            ("bullet", "Genre-based discovery filtered by the user's top 3 preferred genres"),
            ("bullet", "Content similarity from the user's top-8 liked movies (getSimilarMovies + getMovieRecommendations)"),
            ("bullet", "Actor/director-based discovery from the user's top 5 actors and top 3 directors"),
            ("h2", "Stage 2 — Scoring (_scoreMovies)"),
            ("bullet", "Base Score (50%): genre match (40%), actor match (25%), director match (20%), rating (10%), recency (5%)"),
            ("bullet", "MovieEmbeddingService (15%): 64-dimensional vector cosine similarity"),
            ("bullet", "ContextualRecommendationService (15%): time of day, mood, weekday/weekend"),
            ("bullet", "BehaviorTrackingService (10%): swipe speed, detail view duration, revisit count"),
            ("bullet", "CollaborativeFilteringService (5%): co-occurrence matrix with popularity normalization"),
            ("bullet", "DeepLearningService (5%): 18-dimensional feature vector with rule-based fallback (TFLite infrastructure ready)"),
            ("h2", "Stage 3 — Filtering & Deduplication"),
            ("bullet", "Remove already-liked, disliked, watchlisted, and skipped items"),
            ("bullet", "Platform availability filter via StreamingService (TMDB watch/providers API)"),
            ("bullet", "Diversity filter: prevents genre clustering in the visible deck (max 2 genre overlaps in first 10 cards)"),
            ("h2", "Stage 4 — Adaptive Meta-Learning"),
            ("bullet", "AdaptiveWeightingService learns which of the 5 strategies works best per individual user using exponential smoothing"),
            ("bullet", "MatrixFactorizationService maintains 20-dimensional user and movie embedding vectors updated via gradient descent"),
            ("bullet", "OnlineLearningService triggers incremental embedding updates after each interaction (rate-limited to 5-second intervals)"),
        ],
    },
    {
        "heading": "3. RAG — What It Is and How PopMatch Implements the Pattern",
        "level": 1,
        "content": [
            ("h2", "What is RAG?"),
            ("p", "Retrieval-Augmented Generation (RAG) is an AI architecture pattern that solves the core limitations of pure LLMs (stale knowledge, hallucination, missing domain data). Instead of answering from training weights alone, the system:"),
            ("bullet", "RETRIEVES relevant documents/records from an external knowledge base at query time"),
            ("bullet", "AUGMENTS the model's context with that retrieved, up-to-date content"),
            ("bullet", "GENERATES a response grounded in the retrieved information"),
            ("h2", "PopMatch as a RAG System (Without the LLM)"),
            ("p", "PopMatch's recommendation pipeline is structurally identical to a RAG system:"),
            ("table", [
                ["RAG Component", "PopMatch Equivalent"],
                ["Knowledge Base", "TMDB movie catalog + OMDb ratings + social activity feed"],
                ["Retrieval", "4 parallel TMDB query strategies (Stage 1)"],
                ["Augmentation", "User behavioral context injected into scoring (liked movies, swipe speed, time of day)"],
                ["Generation", "_scoreMovies() produces a personalized ranked list"],
                ["Query", "User's implicit query = like/dislike history + mood + current time"],
            ]),
            ("h2", "Real LLM-Based RAG System (Production)"),
            ("p", "When asked about production RAG with LLMs, the steps are:"),
            ("bullet", "Embed user query + documents into the same vector space using a pre-trained model (e.g., text-embedding-3-small, bge-m3, e5-mistral-7b)"),
            ("bullet", "Store document embeddings in a vector DB (Pinecone, Weaviate, Chroma, pgvector)"),
            ("bullet", "At query time: embed the query, run ANN search (FAISS / HNSW) to retrieve top-K docs in milliseconds"),
            ("bullet", "Inject retrieved docs into the LLM prompt as grounding context"),
            ("bullet", "Generate response — grounded, not hallucinated"),
            ("p", "PopMatch's MovieEmbeddingService implements steps 1–3 manually in Dart — the same mathematical principles, just without an LLM at the end. Building it from scratch proves understanding of the underlying math, not just API calls."),
        ],
    },
    {
        "heading": "4. Embeddings — Deep Technical Explanation",
        "level": 1,
        "content": [
            ("h2", "What Are Embeddings?"),
            ("p", "An embedding is a dense, fixed-size numerical vector that encodes the semantic meaning of a data point. The fundamental property: semantically similar items have vectors that are geometrically close in the embedding space, measured by cosine similarity."),
            ("code", "cosine_similarity(A, B) = (A · B) / (||A|| × ||B||)\nRange: [-1, 1], where 1 = identical, 0 = orthogonal, -1 = opposite"),
            ("h2", "PopMatch's 64-Dimensional Movie Embedding Vector (MovieEmbeddingService)"),
            ("table", [
                ["Dimensions", "Count", "Feature", "Encoding"],
                ["0 – 19", "20", "Genre presence", "One-hot per genre (action=1, comedy=0, etc.)"],
                ["20 – 24", "5", "Vote average", "rating / 10.0, repeated 5× for weight"],
                ["25 – 29", "5", "Release year", "(year - 1900) / 130"],
                ["30 – 34", "5", "Popularity", "popularity / 100, clamped to 1.0"],
                ["35 – 39", "5", "Runtime", "runtime / 200"],
                ["40 – 44", "5", "Vote count", "log(voteCount + 1) / 15"],
                ["45 – 49", "5", "Language", "hashCode % 100 / 100 (hash feature)"],
                ["50 – 55", "6", "Text keyword freq.", "Genre keywords in overview text (NLP)"],
                ["56", "1", "Description length", "wordCount / 200"],
                ["57", "1", "Vocabulary diversity", "unique words / total words"],
                ["58", "1", "Positive emotion", "(happy, joy, win, hero) / total words"],
                ["59", "1", "Negative emotion", "(death, dark, fear, evil) / total words"],
                ["60 – 63", "4", "Thematic keywords", "space, future, past, journey prevalence"],
            ]),
            ("p", "After construction: L2-normalize the vector so all vectors have unit length. Cosine similarity then reduces to a simple dot product."),
            ("h2", "Dual Similarity Scoring"),
            ("code", "for each likedMovie in user.likedMovies:\n    sim = cosine_similarity(candidate_vec, liked_vec)\n\nfinal = 0.6 × max_similarity + 0.4 × avg_similarity\nweight = 0.3 + (final × 0.7)   # scaled to [0.3, 1.0]"),
            ("p", "max captures 'does this exactly match something the user loves?' — avg captures 'does this consistently fit their overall taste profile?' The dual approach handles both passionate matches and consistent preferences."),
            ("h2", "Production Scale"),
            ("bullet", "Pre-trained embedding models: text-embedding-ada-002, bge-m3, e5-mistral-7b"),
            ("bullet", "Vector databases: Pinecone, Weaviate, Qdrant, pgvector"),
            ("bullet", "ANN search: FAISS (Facebook), HNSW (Hierarchical Navigable Small World) — billions of vectors in milliseconds"),
            ("bullet", "Building embeddings from scratch (as in PopMatch) proves understanding of the math, not just API usage"),
        ],
    },
    {
        "heading": "5. Machine Learning Training — Six Techniques",
        "level": 1,
        "content": [
            ("h2", "5.1 Item-Based Collaborative Filtering (CollaborativeFilteringService)"),
            ("p", "Every addLikedMovie() call is a real-time training step. The model updates a co-occurrence matrix — a form of online learning, not batch training."),
            ("code", "For every user who liked movie A AND movie B:\n    co_occurrence[A][B] += 1\n\nScore for candidate C given user liked A:\n    raw_score = co_occurrence[A][C]\n    normalized = raw_score / sqrt(movie_popularity[C])\n    # sqrt dampens blockbuster bias — surfaces niche relevant films"),
            ("p", "Square-root popularity normalization prevents popular blockbusters (Avengers, Inception) from dominating co-occurrence counts, giving niche but relevant films a fair chance."),
            ("h2", "5.2 Matrix Factorization — SVD-Like Decomposition (MatrixFactorizationService)"),
            ("p", "Implements a Singular Value Decomposition (SVD)-style approach with 20 latent factors. This is the same technique that powered the Netflix Prize winning algorithm."),
            ("code", "# Latent factors: 20 dimensions\n# Learning rate: 0.01   Regularization: 0.1\n\n# User embedding: userId -> List<double> (20D vector)\n# Movie embedding: movieId -> List<double> (20D vector)\n\n# Prediction:\npredicted_rating = dot(user_vec, movie_vec)  # normalized to [0, 1]\n\n# Gradient descent update:\nerror = actual_rating - predicted_rating\nuser_vec += lr * (error * movie_vec - reg * user_vec)\nmovie_vec += lr * (error * user_vec - reg * movie_vec)"),
            ("p", "Key methods: predict(userId, movieId), updateFromFeedback(), getMovieSimilarity(), getUserSimilarity(), getRecommendedMovies(). Embeddings are persisted to SharedPreferences as JSON."),
            ("h2", "5.3 Adaptive Meta-Learning (AdaptiveWeightingService)"),
            ("p", "The system does not just score movies — it learns which scoring strategies work best for each individual user and adjusts weights continuously via exponential smoothing. This is a multi-armed bandit problem."),
            ("code", "Initial weights per strategy:\n  contentBased = 40%, contextual = 20%, embedding = 20%\n  behavior = 15%, collaborative = 5%\n\nAfter >= 10 feedbacks per strategy:\n  success_rate = likes / (likes + dislikes)\n  # Exponential moving average:\n  new_weight = (old_weight * 0.7) + (target_weight * 0.3)\n  # Re-normalize: all weights must sum to 1.0"),
            ("p", "Context-aware cold/warm start adjustment:"),
            ("bullet", "New users (<5 likes): boost contentBased x1.2 + contextual x1.2, reduce embedding x0.8"),
            ("bullet", "Experienced users (>20 likes): boost embedding x1.3 + collaborative x1.5, reduce contentBased x0.9"),
            ("bullet", "Active session: boost behavior weight x1.2"),
            ("h2", "5.4 Online Learning (OnlineLearningService)"),
            ("p", "Incremental learning from every user interaction (like, dislike, skip, view). Rate-limited to 5-second minimum intervals to avoid excessive updates. Ratings: like=+1, dislike=-0.5, neutral=0. Triggers MatrixFactorization gradient descent updates and updates embeddings incrementally when preferences change. Stores update history in SharedPreferences for model freshness tracking."),
            ("h2", "5.5 Implicit Behavioral Feedback (BehaviorTrackingService)"),
            ("p", "Converts interaction patterns into continuous interest scores — the same technique used by Netflix and Spotify. Implicit feedback is often more predictive than explicit ratings."),
            ("code", "Interest score for a movie =\n    (detailViewCount * 0.3)\n  + clamp(avgTimeOnDetailPage / 10.0, 0, 2.0)\n  + (revisitCount * 0.5)\n  + 0.2 if swipeSpeed > 2000ms  # hesitation = interest\n  + explicitSignal              # +1.0 like, -0.5 dislike\n\nbehavior_weight = 0.5 + (normalized_interest * 0.5)  # [0.5, 1.0]\nbehavior_weight = 0.5 - (normalized_disinterest * 0.5)  # [0.0, 0.5]"),
            ("h2", "5.6 Contextual Multi-Armed Bandits (ContextualRecommendationService)"),
            ("p", "Adjusts scoring weights based on observable context — time of day, day of week, and current mood. The reward function changes with context rather than being stationary."),
            ("code", "Morning  (6-12):  comedy x1.2, animation x1.15, horror x0.8\nEvening (17-22): action x1.15, drama x1.1, thriller x1.1\nNight   (22-6):  horror x1.2, thriller x1.15, animation x0.85\n\nWeekend: prefer longer movies (>120 min) x1.1\nWeekday: prefer shorter movies (<100 min) x1.05\n\nMood multiplier: 1.0 + (matchingGenres * 0.15)"),
            ("h2", "5.7 Deep Learning Infrastructure (DeepLearningService)"),
            ("p", "TensorFlow Lite infrastructure deployed, model not yet trained. Uses rule-based fallback scoring. Feature vector ready for neural network training."),
            ("code", "18-dimensional feature vector:\n[0]   user.likedMovieCount (normalized)\n[1]   user.dislikedMovieCount (normalized)\n[2]   movie.voteAverage / 10.0\n[3]   movie.popularity / 100.0\n[4]   movie.voteCount / 1000.0\n[5-14] genre one-hot (top 10 TMDB genre IDs)\n[15]  (releaseYear - 1900) / 130.0\n[16]  runtime / 200.0\n[17]  0.0 (reserved)\n\nNext step: collect labeled swipe data (right=1, left=0),\ntrain 3-4 layer feedforward network, export to .tflite"),
        ],
    },
    {
        "heading": "6. AI Agents — Framing the Architecture",
        "level": 1,
        "content": [
            ("h2", "What Is an AI Agent?"),
            ("p", "An AI agent is a system that: Perceives its environment (input) → Reasons about what to do (planning/decision) → Acts (calls tools/APIs) → Observes the outcome → Learns and updates its state."),
            ("h2", "PopMatch's Recommendation System as an AI Agent"),
            ("table", [
                ["Agent Property", "PopMatch Implementation"],
                ["Perception", "User swipes, view times, time of day, liked movie history, current mood"],
                ["Reasoning", "_scoreMovies() — multi-factor weighted evaluation of 100+ candidates"],
                ["Tool Use", "TMDB API, OMDb API, StreamingService — external tools called dynamically"],
                ["Action", "Surfaces ranked movie list to the swipe interface"],
                ["Observation", "Records outcome: liked? disliked? opened detail? time spent?"],
                ["Learning", "AdaptiveWeightingService + OnlineLearningService update models from outcomes"],
                ["Short-term Memory", "BehaviorTrackingService in-memory session signals"],
                ["Long-term Memory", "SharedPreferences: co-occurrence matrix, MF embeddings, strategy weights"],
            ]),
            ("h2", "Production LLM-Based Agent (Like Microsoft Copilot)"),
            ("p", "The same pattern implemented with LLMs using the ReAct (Reason + Act) loop:"),
            ("code", "# ReAct agent loop\nwhile not done:\n    thought = llm.think(system_prompt, history, current_observation)\n    action = llm.choose_tool(thought, available_tools)\n    observation = tool_registry[action.name].run(action.params)\n    history.append((thought, action, observation))\n    done = llm.should_stop(history)"),
            ("p", "Key LLM agent concepts: Tool use / Function calling, Chain-of-thought planning, Short-term memory (conversation context) + Long-term memory (vector DB), Orchestration frameworks: LangChain, LlamaIndex, Microsoft Semantic Kernel (for Copilot), CrewAI, AutoGen."),
        ],
    },
    {
        "heading": "7. Data Architecture — Full Pipeline",
        "level": 1,
        "content": [
            ("h2", "User Model — What Is Stored"),
            ("p", "The User model carries all preference and interaction state used to drive personalization:"),
            ("bullet", "Authentication: id (Firebase UID), email, displayName, photoURL"),
            ("bullet", "Movie tracking: likedMovies, dislikedMovies, watchlist — the primary training signals"),
            ("bullet", "Show tracking: likedShows, dislikedShows, watchlistShows"),
            ("bullet", "Context: currentMood (string) — drives ContextualRecommendationService"),
            ("bullet", "Preferences map: onboardingCompleted (bool), selectedGenres (List<int>), selectedPlatforms (List<String>)"),
            ("h2", "Onboarding → Recommendation Integration"),
            ("p", "The 3-step onboarding (welcome → genre selection → platform selection) seeds the cold-start recommendation engine:"),
            ("bullet", "Genres selected during onboarding become the initial preference signal"),
            ("bullet", "UserPreferenceAnalyzer uses onboarding genres as defaults when liked movie count is below threshold"),
            ("bullet", "Blend weight (0.0 → 1.0) progressively personalizes recommendations as likes accumulate (0 likes = pure onboarding, 3+ likes = full personalization)"),
            ("bullet", "UserPreferenceAnalyzer extracts top genres, actors, directors, and rating ranges from up to 30 liked movies in parallel"),
            ("bullet", "Rating range: 25th/75th percentile of liked movie ratings (ignores outliers) — defines the user's quality bar"),
            ("h2", "Performance Optimizations"),
            ("bullet", "Infinite swipe buffer: 30-movie minimum pre-loaded before they are needed — zero wait on swipe"),
            ("bullet", "Staged loading: instant load from disk cache → fresh load in background → merge seamlessly"),
            ("bullet", "_pendingMovies queue: background-loaded movies held until main deck empties — no mid-swipe pop-in"),
            ("bullet", "Session cache: UserPreferencesSessionCache avoids reanalyzing user preferences on every tab switch"),
            ("bullet", "MovieCacheService: 24-hour TTL, 50-movie LRU cache for detail screens — preloads before navigation"),
            ("bullet", "Rate limiting: 850ms–900ms throttle on preload kicks to avoid API flooding"),
            ("bullet", "Deferred removal: card removed from deck 4 seconds after swipe — allows undo without flickering"),
            ("h2", "Complete Data Flow"),
            ("code", "TMDB API\n  -> TMDBService (8s timeout, test mode support)\n  -> Movie.fromJson() deserialization\n  -> MovieProvider._scoreMovies()\n       |- UserPreferencesSessionCache (top genres, actors, directors)\n       |- Base scoring (genre/actor/director/rating/recency)\n       |- MovieEmbeddingService (64D cosine similarity)\n       |- ContextualRecommendationService (time/mood)\n       |- BehaviorTrackingService (implicit signals)\n       |- CollaborativeFilteringService (co-occurrence)\n       |- MatrixFactorizationService (SVD-like)\n  -> _applyDiversityFilter()\n  -> StreamingService platform filter\n  -> _filteredMovies (visible swipe deck)\n  -> RetroCinemaMovieCard (UI)\n  -> User swipes\n  -> AuthProvider (like/dislike/watchlist recorded)\n  -> OnlineLearningService (incremental MF update)\n  -> BehaviorTrackingService (implicit signal update)\n  -> AdaptiveWeightingService (strategy weight update)\n  -> OMDb API (background: IMDb/RT/Metacritic ratings)"),
        ],
    },
    {
        "heading": "8. Social Architecture & Governance",
        "level": 1,
        "content": [
            ("h2", "Social Features (SocialService + Cloud Functions)"),
            ("p", "Social features run entirely through Firebase Cloud Functions (Node.js, region: us-central1) with Firestore as the data store. The client-side SocialService gracefully degrades when Firebase is disabled."),
            ("bullet", "Follow system: sendFollowRequest(), respondToFollowRequest(), unfollow()"),
            ("bullet", "User discovery: searchUsers() (case-insensitive, by name/email), getSuggestedUsers() (Jaccard similarity — minimum 2 shared liked movies)"),
            ("bullet", "Social activity feed: recordActivity() respects per-user privacy settings server-side before writing to Firestore"),
            ("bullet", "Friends feed: getFriendsFeed() filters out activities marked private — privacy enforced at Cloud Function level, not client"),
            ("h2", "Governance Layers"),
            ("p", "Governance in PopMatch runs at three levels:"),
            ("bullet", "Algorithm level: A/B testing with 3 persistent variants (variantA 50%, variantB 30%, variantC 20%). Tracks Precision@K, Recall@K, NDCG@K, Diversity, Novelty, and Coverage per variant"),
            ("bullet", "Privacy level: Social activity sharing enforced server-side in Cloud Functions — client cannot bypass visibility rules"),
            ("bullet", "Observability level: Every recommendation carries a recommendationStrategy field identifying which algorithm surfaced it — explainability built in"),
            ("bullet", "Feature flags: Compile-time dart-define flags (SOCIAL_UI_ENABLED, FRIENDS_FEED_ENABLED) for controlled feature rollout"),
            ("h2", "Recommendation Metrics (RecommendationMetricsService)"),
            ("table", [
                ["Metric", "Formula", "Weight"],
                ["Precision@K", "Liked items in top K / K", "30%"],
                ["Recall@K", "Liked items in top K / total liked", "25%"],
                ["NDCG@K", "DCG / IDCG (log-discounted rank quality)", "25%"],
                ["Diversity", "Genre spread across recommendations", "10%"],
                ["Novelty", "Inverse popularity (surfacing non-obvious items)", "5%"],
                ["Coverage", "Fraction of catalog that can be recommended", "5%"],
            ]),
            ("p", "Overall quality score = Precision(0.30) + Recall(0.25) + NDCG(0.25) + Diversity(0.10) + Novelty(0.05) + Coverage(0.05). Stored in SharedPreferences (last 1000 evaluations) with timestamps for temporal analysis."),
        ],
    },
    {
        "heading": "9. Python Backend Mapping",
        "level": 1,
        "content": [
            ("p", "The job requires Python backend development. Every PopMatch component maps directly to Python equivalents:"),
            ("table", [
                ["PopMatch (Dart)", "Python Equivalent"],
                ["TMDBService", "FastAPI + httpx async client"],
                ["MovieEmbeddingService", "numpy + sklearn.preprocessing.normalize + cosine_similarity"],
                ["CollaborativeFilteringService", "scipy.sparse lil_matrix"],
                ["MatrixFactorizationService", "numpy gradient descent or surprise library (SVD)"],
                ["Future.wait() parallel calls", "asyncio.gather() or ThreadPoolExecutor"],
                ["SharedPreferences ML state", "Redis (with TTL) or PostgreSQL"],
                ["Background preload timer", "Celery async tasks + Redis broker"],
                ["AdaptiveWeightingService", "scipy.stats + scikit-learn"],
                ["A/B significance test", "scipy.stats.ttest_ind (p_value < 0.05)"],
                ["Vector DB (production)", "Pinecone / pgvector / Weaviate"],
                ["Cloud Functions (Node.js)", "FastAPI microservices or AWS Lambda (Python)"],
            ]),
            ("h2", "Python Recommendation API Example"),
            ("code", "@app.post('/recommendations')\nasync def get_recommendations(user: UserProfile, limit: int = 30):\n    # Stage 1: Parallel retrieval\n    candidates = await asyncio.gather(\n        tmdb.get_trending(),\n        tmdb.discover_by_genres(user.top_genres),\n        tmdb.get_similar(user.top_liked_ids[:8])\n    )\n    all_candidates = deduplicate(flatten(candidates))\n\n    # Stage 2: Multi-factor scoring\n    for movie in all_candidates:\n        base  = compute_base_score(movie, user)      # 50%\n        embed = embedding_svc.get_weight(movie, user) # 15%\n        ctx   = contextual_svc.get_weight(movie)      # 15%\n        behav = behavior_svc.get_weight(movie.id)     # 10%\n        collab = cf_svc.get_score(movie.id, user)     # 5%\n        score = 0.50*base + 0.15*embed + 0.15*ctx \\\n              + 0.10*behav + 0.05*collab\n\n    # Stage 3: Diversity filter + return\n    return diversity_filter(sorted(scored, reverse=True))[:limit]"),
            ("h2", "Python Embedding Service Example"),
            ("code", "import numpy as np\nfrom sklearn.preprocessing import normalize\nfrom sklearn.metrics.pairwise import cosine_similarity\n\ndef build_embedding(movie: Movie) -> np.ndarray:\n    vec = np.zeros(64)\n    for i, gid in enumerate(movie.genre_ids[:20]):\n        if gid in GENRE_INDEX: vec[GENRE_INDEX[gid]] = 1.0\n    vec[20:25] = movie.vote_average / 10.0\n    vec[25:30] = (movie.release_year - 1900) / 130.0\n    # ... remaining dimensions\n    return normalize(vec.reshape(1, -1))[0]  # L2 normalize\n\ndef get_weight(candidate, liked_movies) -> float:\n    c_vec = build_embedding(candidate)\n    l_vecs = np.array([build_embedding(m) for m in liked_movies])\n    sims = cosine_similarity([c_vec], l_vecs)[0]\n    return 0.6 * sims.max() + 0.4 * sims.mean()"),
        ],
    },
    {
        "heading": "10. LLM Orchestration, Performance, and Governance",
        "level": 1,
        "content": [
            ("h2", "Orchestration Patterns"),
            ("bullet", "LangChain / LlamaIndex: Pipeline-based orchestration for RAG + agents"),
            ("bullet", "Microsoft Semantic Kernel: Copilot-specific framework — Planner + Skills + Memory + Connectors"),
            ("bullet", "CrewAI / AutoGen: Multi-agent systems where specialized sub-agents collaborate"),
            ("bullet", "Prefect / Airflow: Orchestrating ML pipelines and data workflows in production"),
            ("h2", "Performance Patterns (from PopMatch)"),
            ("bullet", "Parallelism: 4 discovery strategies run simultaneously (Future.wait in Dart, asyncio.gather in Python)"),
            ("bullet", "Fallback chains: personalized fails → curated starter movies. In production: define retry/fallback at the orchestration layer"),
            ("bullet", "Rate limiting: 850ms throttle on preload kicks. In production: token bucket or leaky bucket algorithm"),
            ("bullet", "Multi-level caching: session cache (in-memory) + disk cache (SharedPreferences) + TTL-based expiry"),
            ("h2", "Production AI Governance (Beyond PopMatch)"),
            ("bullet", "Model cards: document training data, evaluation metrics, known failure modes, intended use"),
            ("bullet", "Bias detection: monitor recommendation distribution across demographic groups"),
            ("bullet", "Explainability: the recommendationStrategy field (already in PopMatch's Movie model) explains why each item was surfaced"),
            ("bullet", "PII handling: never store raw behavioral data with user PII in the same store"),
            ("bullet", "Drift detection: compare current recommendation distributions against baseline — alert on significant shifts"),
            ("bullet", "Human-in-the-loop: escalation paths for low-confidence agent decisions"),
        ],
    },
    {
        "heading": "11. Interview Talking Points",
        "level": 1,
        "content": [
            ("h2", "On AI Agents"),
            ("p", "\"In PopMatch, I built a recommendation agent that follows the classic perceive-reason-act-learn loop. The orchestrator coordinates 7 specialized modules, each providing a different signal. This is architecturally identical to tool-using LLM agents — the agent decides which tools to call (TMDB discover, OMDb ratings, streaming availability), collects results, and synthesizes a ranked output. In a Copilot-based system I would add natural language planning using Semantic Kernel's Planner before tool invocation.\""),
            ("h2", "On RAGs"),
            ("p", "\"PopMatch implements the retrieval-augment-generate pattern without an LLM. The retrieval stage runs 4 parallel TMDB queries. The augmentation injects behavioral context — liked movies, swipe speed, time of day. The generation is the scoring function that produces a ranked list. In an LLM-based RAG, you replace the scoring function with a language model reasoning over retrieved documents. I understand both the indexing side (building embeddings, storing in vector DBs) and the query side (ANN search, re-ranking, context window management).\""),
            ("h2", "On Embeddings"),
            ("p", "\"I hand-built 64-dimensional movie embeddings encoding genre, rating, year, popularity, runtime, and NLP text features from movie descriptions. I implemented L2 normalization and cosine similarity from scratch with a dual max+average approach that captures both passionate matches and consistent taste profiles. In production I would use a pre-trained model and a vector DB, but building it manually means I fully understand the underlying math.\""),
            ("h2", "On ML Training"),
            ("p", "\"I implemented five forms of ML training: (1) online collaborative filtering — every like updates a co-occurrence matrix in real-time; (2) matrix factorization — SVD-like 20-dimensional latent factor model updated via gradient descent; (3) adaptive meta-learning — exponential smoothing learns optimal strategy weights per user; (4) implicit feedback — swipe speed, detail page duration, and revisit count become continuous interest scores; (5) contextual bandits — genre weights shift based on time of day and mood.\""),
            ("h2", "On Governance"),
            ("p", "\"Governance runs at three levels: (1) algorithm — A/B testing 3 variants tracking precision, recall, NDCG, diversity, novelty, and coverage; (2) privacy — social activity visibility enforced server-side in Cloud Functions, clients cannot bypass; (3) observability — every recommendation carries a recommendationStrategy field for explainability. In production I would add model cards, drift detection, and bias auditing.\""),
            ("h2", "On Python Backend"),
            ("p", "\"All ML logic in PopMatch is in Dart. I can map every component to Python: FastAPI for the recommendation API, httpx for async TMDB calls, numpy/sklearn for embeddings, scipy for collaborative filtering, Redis for caching, Celery for background pre-computation, and pgvector or Pinecone for embedding storage. The architecture is identical — only the language changes.\""),
        ],
    },
    {
        "heading": "12. Common Interview Questions & Answers",
        "level": 1,
        "content": [
            ("h2", "Q: What is RAG and when would you use it?"),
            ("p", "Retrieval-Augmented Generation retrieves relevant documents from an external knowledge base at query time to ground LLM responses. Use it when: model training data is stale, domain-specific knowledge is required, or hallucination is unacceptable. Requires: embedding model + vector DB + re-ranker + LLM. PopMatch is structurally a RAG system — the scoring function plays the LLM's role."),
            ("h2", "Q: How do embeddings work?"),
            ("p", "Dense numerical vectors where similar items cluster together in vector space. Built by encoding semantic features into fixed-size arrays and L2-normalizing. Similarity measured by cosine distance (dot product after normalization). In PopMatch: 64D manual embeddings. In production: pre-trained models (ada-002, bge-m3) stored in Pinecone/Weaviate, searched with FAISS/HNSW."),
            ("h2", "Q: Collaborative vs content-based filtering?"),
            ("p", "Content-based: compare item attributes (genres, cast, description) — 'this movie shares your liked genres.' Collaborative: compare user behavior — 'users who liked what you liked also liked this.' PopMatch implements both and combines them with adaptive weights. Cold-start problem: collaborative needs data, content-based works immediately."),
            ("h2", "Q: How do you handle cold start?"),
            ("p", "Below 3 likes: curated popular + onboarding genre-based recommendations (content-based only, no collaborative signal). Above 3 likes: blend weight increases progressively to 1.0 — full personalized pipeline with all 7 services. For LLM agents: onboarding prompts build an explicit preference prior before behavioral data exists."),
            ("h2", "Q: What is NDCG?"),
            ("p", "Normalized Discounted Cumulative Gain — measures ranking quality. Items at position 1 contribute more than position 10 (logarithmic discount: 1/log2(rank+1)). Normalized means relative to the ideal ranking (IDCG). Formula: NDCG@K = DCG@K / IDCG@K. Used in PopMatch's RecommendationMetricsService with 25% weight in overall quality score."),
            ("h2", "Q: What is matrix factorization?"),
            ("p", "Decomposes the user-item interaction matrix into two low-dimensional matrices: user embeddings and item embeddings. The dot product of a user vector and item vector predicts the user's rating. SVD (Singular Value Decomposition) is the canonical approach — powered the Netflix Prize winning algorithm. PopMatch implements this with 20 latent factors, learning rate 0.01, and L2 regularization 0.1."),
            ("h2", "Q: What is a multi-armed bandit?"),
            ("p", "An online learning problem where an agent chooses between K strategies (arms), observes a reward, and learns which arms yield the highest reward over time. In PopMatch: each of the 5 scoring strategies is an arm. AdaptiveWeightingService is the bandit — it uses exponential smoothing (not pure exploration/exploitation) to converge on optimal per-user weights."),
            ("h2", "Q: How would you scale this to production in Python?"),
            ("p", "FastAPI microservice for the recommendation API, async httpx for TMDB, Redis for caching embeddings and scores (TTL 24h), Celery for background pre-computation triggered by user events, PostgreSQL + pgvector for storing movie embeddings, Pinecone or Weaviate for ANN search at scale, Prometheus + Grafana for metrics, Optimizely or custom for A/B test management, MLflow for model tracking and versioning."),
        ],
    },
    {
        "heading": "13. Critical Files to Review",
        "level": 1,
        "content": [
            ("table", [
                ["File", "Why It Matters"],
                ["lib/services/movie_embedding_service.dart", "Core embeddings — know all 64 dimensions and dual scoring"],
                ["lib/services/collaborative_filtering_service.dart", "Co-occurrence matrix, sqrt popularity normalization"],
                ["lib/services/matrix_factorization_service.dart", "SVD-like 20D latent factors, gradient descent updates"],
                ["lib/services/online_learning_service.dart", "Incremental learning loop, rate limiting"],
                ["lib/services/adaptive_weighting_service.dart", "Meta-learning, exponential smoothing, multi-armed bandit"],
                ["lib/services/behavior_tracking_service.dart", "Implicit feedback signals, interest score formula"],
                ["lib/services/contextual_recommendation_service.dart", "Circadian patterns, mood multipliers"],
                ["lib/providers/movie_provider.dart", "Full orchestration pipeline — the agent loop"],
                ["lib/services/ab_testing_service.dart", "Variant assignment, metrics, statistical comparison"],
                ["lib/services/recommendation_metrics_service.dart", "Precision, Recall, NDCG formulas"],
                ["lib/services/user_preference_analyzer.dart", "Preference extraction, blend weight, cold start"],
                ["functions/index.js", "Cloud Functions, social governance, privacy enforcement"],
                ["docs/RECOMMENDATION_ALGORITHM.md", "Exact weight breakdown for the full scoring pipeline"],
            ]),
        ],
    },
]


# ─── DOCX Generator ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_code_block(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    # Light grey background via paragraph shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)


def add_table(doc, rows):
    headers = rows[0]
    table = doc.add_table(rows=len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(cell, "1F3864")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(rows[1:], 1):
        bg = "EBF5FB" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(cell, bg)
    doc.add_paragraph()


def generate_docx(path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("PopMatch — AI/ML Interview Preparation Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    subtitle = doc.add_paragraph("Architecture · RAG · Embeddings · ML Training · AI Agents · Python Backend")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    for section in SECTIONS:
        level = section["level"]
        heading = doc.add_heading(section["heading"], level=level)
        heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        for item_type, item_content in section["content"]:
            if item_type == "p":
                p = doc.add_paragraph(item_content)
                p.runs[0].font.size = Pt(10)
            elif item_type == "h2":
                h = doc.add_heading(item_content, level=2)
                h.runs[0].font.color.rgb = RGBColor(0x1A, 0x5E, 0x8A)
            elif item_type == "bullet":
                p = doc.add_paragraph(item_content, style="List Bullet")
                p.runs[0].font.size = Pt(10)
            elif item_type == "code":
                add_code_block(doc, item_content)
            elif item_type == "table":
                add_table(doc, item_content)

    doc.save(path)
    print(f"DOCX saved: {path}")


# ─── PDF Generator ────────────────────────────────────────────────────────────

_CHAR_MAP = [
    (chr(0x2014), "-"),   # em dash
    (chr(0x2013), "-"),   # en dash
    (chr(0x00d7), "x"),   # multiplication sign
    (chr(0x00b7), "."),   # middle dot
    (chr(0x2019), "'"),   # right single quotation
    (chr(0x2018), "'"),   # left single quotation
    (chr(0x201c), '"'),   # left double quotation
    (chr(0x201d), '"'),   # right double quotation
    (chr(0x2022), "-"),   # bullet
    (chr(0x2026), "..."), # ellipsis
    (chr(0x2192), "->"),  # right arrow
    (chr(0x2190), "<-"),  # left arrow
    (chr(0x2194), "<->"), # left-right arrow
    (chr(0x2713), "OK"),  # check mark
    (chr(0x00b2), "2"),   # superscript 2
    (chr(0x00b3), "3"),   # superscript 3
    (chr(0x00b1), "+/-"), # plus-minus
    (chr(0x2264), "<="),  # less than or equal
    (chr(0x2265), ">="),  # greater than or equal
    (chr(0x2260), "!="),  # not equal
    (chr(0x03b1), "alpha"), # greek alpha
    (chr(0x03b2), "beta"),  # greek beta
]


def sanitize(text):
    for char, replacement in _CHAR_MAP:
        text = text.replace(char, replacement)
    return text.encode("latin-1", errors="replace").decode("latin-1")


class PDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(100, 100, 100)
        self.cell(0, 8, "PopMatch - AI/ML Interview Preparation Guide", align="C")
        self.ln(2)
        self.set_draw_color(31, 56, 100)
        self.set_line_width(0.3)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(4)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")

    def chapter_title(self, text, level=1):
        text = sanitize(text)
        if level == 1:
            self.set_font("Helvetica", "B", 13)
            self.set_fill_color(31, 56, 100)
            self.set_text_color(255, 255, 255)
            self.cell(0, 8, text, new_x="LMARGIN", new_y="NEXT", fill=True)
            self.ln(2)
        else:
            self.set_font("Helvetica", "B", 11)
            self.set_text_color(26, 94, 138)
            self.cell(0, 7, text, new_x="LMARGIN", new_y="NEXT")
            self.ln(1)
        self.set_text_color(0, 0, 0)

    def body_text(self, text):
        self.set_font("Helvetica", "", 9.5)
        self.set_text_color(30, 30, 30)
        self.multi_cell(0, 5, sanitize(text))
        self.ln(2)

    def bullet_item(self, text):
        self.set_font("Helvetica", "", 9.5)
        self.set_text_color(30, 30, 30)
        self.set_x(self.get_x() + 5)
        self.multi_cell(0, 5, f"-  {sanitize(text)}", new_x="LMARGIN", new_y="NEXT")

    def code_block(self, text):
        self.set_fill_color(245, 245, 245)
        self.set_draw_color(200, 200, 200)
        self.set_font("Courier", "", 7.5)
        self.set_text_color(46, 125, 50)
        self.set_x(self.get_x() + 5)
        self.multi_cell(
            w=self.epw - 5,
            h=4,
            text=sanitize(text),
            border=1,
            fill=True,
            new_x="LMARGIN",
            new_y="NEXT",
        )
        self.ln(2)
        self.set_text_color(0, 0, 0)

    def draw_table(self, rows):
        headers = rows[0]
        col_w = (self.epw) / len(headers)
        # Header row
        self.set_fill_color(31, 56, 100)
        self.set_text_color(255, 255, 255)
        self.set_font("Helvetica", "B", 8)
        for h in headers:
            self.cell(col_w, 6, sanitize(h)[:40], border=1, fill=True)
        self.ln()
        # Data rows
        self.set_text_color(30, 30, 30)
        for i, row in enumerate(rows[1:]):
            if i % 2 == 0:
                self.set_fill_color(235, 245, 251)
            else:
                self.set_fill_color(255, 255, 255)
            self.set_font("Helvetica", "", 7.5)
            lines_needed = 1
            for val in row:
                chars_per_line = int(col_w / 2.1)
                lines_needed = max(lines_needed, max(1, len(val) // chars_per_line + 1))
            row_h = max(6, lines_needed * 4)
            for val in row:
                self.multi_cell(col_w, row_h / lines_needed, sanitize(val), border=1, fill=True, new_x="RIGHT", new_y="LAST", max_line_height=4)
            self.ln(row_h)
        self.ln(3)
        self.set_text_color(0, 0, 0)


def generate_pdf(path):
    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Title page
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(31, 56, 100)
    pdf.ln(10)
    pdf.cell(0, 12, "PopMatch", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 8, "AI/ML Interview Preparation Guide", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 6, "Architecture | RAG | Embeddings | ML Training | AI Agents | Python Backend", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)
    pdf.set_text_color(0, 0, 0)

    for section in SECTIONS:
        level = section["level"]
        pdf.chapter_title(section["heading"], level=level)

        for item_type, item_content in section["content"]:
            if item_type == "p":
                pdf.body_text(item_content)
            elif item_type == "h2":
                pdf.chapter_title(item_content, level=2)
            elif item_type == "bullet":
                pdf.bullet_item(item_content)
            elif item_type == "code":
                pdf.code_block(item_content)
            elif item_type == "table":
                pdf.draw_table(item_content)

    pdf.output(path)
    print(f"PDF saved: {path}")


# ─── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    docx_path = os.path.join(OUTPUT_DIR, "PopMatch_Interview_Guide.docx")
    pdf_path = os.path.join(OUTPUT_DIR, "PopMatch_Interview_Guide.pdf")
    generate_docx(docx_path)
    generate_pdf(pdf_path)
    print("\nDone! Both files saved to docs/")
